import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy
import random
from io import BytesIO
import re

st.set_page_config(page_title="منصة الامتحانات", layout="centered")
st.title("نظام توليد الأسئلة الامتحانية (4 خيارات - قالب مفصول)")

# تأكد من وضع اسم القالب الصحيح هنا
TEMPLATE_FILE = 'template.docx'
# --- دالة المحاذاة ---
def force_rtl(paragraph):
    """تجبر الفقرة على المحاذاة لليمين"""
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.right_indent = None
    paragraph.paragraph_format.first_line_indent = None
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.bidi = True
    if paragraph.runs:
        for run in paragraph.runs:
            run.font.rtl = True

# --- دوال التوسيع ---
def add_row_copy(table, row_idx):
    if row_idx < 0 or row_idx >= len(table.rows): return
    row_copy = copy.deepcopy(table.rows[row_idx]._tr)
    table._tbl.append(row_copy)

def expand_mcq_table(table, current_slots, target_slots):
    needed = target_slots - current_slots
    if needed > 0:
        last_q_row_idx = len(table.rows) - 2
        last_opt_row_idx = len(table.rows) - 1
        for _ in range(needed):
            add_row_copy(table, last_q_row_idx)
            add_row_copy(table, last_opt_row_idx)

def expand_tf_table(table, current_slots, target_slots):
    needed = target_slots - current_slots
    if needed > 0:
        last_row_idx = len(table.rows) - 1
        for _ in range(needed):
            add_row_copy(table, last_row_idx)

def set_document_font_size(doc, size):
    for p in doc.paragraphs:
        for run in p.runs: run.font.size = Pt(size)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs: run.font.size = Pt(size)

def is_header_table(table):
    txt = "".join(cell.text for row in table.rows for cell in row.cells)
    return ("جامعة" in txt or "الامتحان" in txt)

# --- قراءة الأسئلة (4 خيارات) ---
def read_questions(file):
    doc = Document(file)
    mcq_list = []
    tf_list = []
    current_mode = None
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    i = 0
    while i < len(lines):
        line = lines[i]
        if "# اختياري" in line:
            current_mode = "MCQ"; i += 1; continue
        elif "# صح وخطأ" in line:
            current_mode = "TF"; i += 1; continue
            
        if current_mode == "MCQ":
            if i + 4 < len(lines):
                q = lines[i]
                opts = lines[i+1:i+5] 
                if not any("#" in opt for opt in opts):
                    mcq_list.append({"q": q, "opts": opts})
                    i += 5; continue
        elif current_mode == "TF":
            tf_list.append(line)
            i += 1; continue
        i += 1
    return mcq_list, tf_list

# --- توليد الامتحان ---
def generate_exam(mcq_data, tf_data, template_path, target_mcq_count, target_tf_count):
    doc = Document(template_path)
    
    random.shuffle(mcq_data)
    random.shuffle(tf_data)
    
    final_mcq = mcq_data[:target_mcq_count]
    final_tf = tf_data[:target_tf_count]
    
    mcq_idx = 0
    tf_idx = 0
    current_shuffled_opts = None 
    
    # === التوسيع ===
    for table in doc.tables:
        if is_header_table(table): continue
        row_txt = ""
        try:
            for row in table.rows[:3]: 
                row_txt += "".join([c.text for c in row.cells])
        except: pass

        if "A" in row_txt and "B" in row_txt:
            current_slots = sum(1 for row in table.rows if "A" in "".join([c.text for c in row.cells]))
            if target_mcq_count > current_slots:
                expand_mcq_table(table, current_slots, target_mcq_count)
        elif "(" in row_txt and ")" in row_txt and "A" not in row_txt:
            current_slots = sum(1 for row in table.rows if "(" in "".join([c.text for c in row.cells]))
            if target_tf_count > current_slots:
                expand_tf_table(table, current_slots, target_tf_count)

    # === التعبئة والمحاذاة ===
    for table in doc.tables:
        if is_header_table(table): continue

        row_txt_sample = ""
        try:
            for row in table.rows[:3]: 
                row_txt_sample += "".join([c.text for c in row.cells])
        except: pass

        if "A" in row_txt_sample and "B" in row_txt_sample:
            for row in table.rows:
                cells = row.cells
                full_row = "".join([c.text for c in cells])
                
                # أ) سؤال الاختياري (تم التعديل هنا)
                if "A" not in full_row and bool(re.search(r'\d+-', full_row)):
                    if mcq_idx < len(final_mcq):
                        current_opts = final_mcq[mcq_idx]['opts']
                        current_shuffled_opts = list(current_opts)
                        random.shuffle(current_shuffled_opts)
                        q_text = final_mcq[mcq_idx]['q']
                        
                        for i in range(len(cells)):
                            # إذا وجدنا خلية الترقيم مثل 1-
                            if re.search(r'\d+-', cells[i].text):
                                # نضع نص السؤال في الخلية التي تليها مباشرة
                                if i + 1 < len(cells):
                                    cells[i+1].text = q_text
                                    for p in cells[i+1].paragraphs:
                                        force_rtl(p)
                                else:
                                    # حالة احتياطية إذا لم توجد خلية تالية
                                    cells[i].text = cells[i].text.strip() + " " + q_text
                                    for p in cells[i].paragraphs:
                                        force_rtl(p)
                                break # نوقف البحث في هذا الصف بمجرد وضع السؤال
                
                # ب) خيارات الاختياري
                elif "A" in full_row and current_shuffled_opts:
                    opt_map = {
                        'A': current_shuffled_opts[0], 
                        'B': current_shuffled_opts[1], 
                        'C': current_shuffled_opts[2], 
                        'D': current_shuffled_opts[3]
                    }
                    for i in range(len(cells)):
                        ct = cells[i].text.strip().replace(",", "")
                        if ct in opt_map and i+1 < len(cells):
                            target_cell = cells[i+1]
                            target_cell.text = opt_map[ct]
                            for p in target_cell.paragraphs:
                                force_rtl(p)
                    mcq_idx += 1
                    current_shuffled_opts = None

        elif "(" in row_txt_sample and ")" in row_txt_sample and "A" not in row_txt_sample:
            for row in table.rows:
                if tf_idx < len(final_tf):
                    full_row = "".join([c.text for c in row.cells])
                    if "..." in full_row and "(" in full_row:
                         for cell in row.cells:
                            for p in cell.paragraphs:
                                if "..." in p.text:
                                    p.text = re.sub(r'\.{3,}', final_tf[tf_idx], p.text)
                                    force_rtl(p)
                         tf_idx += 1

    set_document_font_size(doc, 10)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- الواجهة ---
st.sidebar.header("لوحة التحكم")
uploaded_file = st.file_uploader("1. ارفع ملف بنك الأسئلة", type=['docx'])

if uploaded_file:
    all_mcq, all_tf = read_questions(uploaded_file)
    if not all_mcq and not all_tf:
        st.error("لا توجد أسئلة، يرجى التأكد من تنسيق بنك الأسئلة.")
    else:
        st.success(f"المتوفر: {len(all_mcq)} سؤال اختياري، {len(all_tf)} سؤال صح وخطأ.")
        st.markdown("---")
        c1, c2 = st.columns(2)
        with c1:
            mcq_count = st.number_input("عدد الاختيارات", 0, len(all_mcq), min(20, len(all_mcq)))
        with c2:
            tf_count = st.number_input("عدد الصح والخطأ", 0, len(all_tf), min(10, len(all_tf)))
            
        if st.button("توليد الامتحان"):
            try:
                final_file = generate_exam(all_mcq, all_tf, TEMPLATE_FILE, mcq_count, tf_count)
                st.download_button("📥 تحميل الامتحان المولد", final_file, "Exam_4_Options_Fixed.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                st.balloons()
                st.success("تم توليد الامتحان بنجاح!")
            except Exception as e:
                st.error(f"حدث خطأ أثناء التوليد: {e}")
